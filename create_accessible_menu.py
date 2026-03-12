"""
Accessible Microsoft Word Menu Creator for Akron Family Restaurant
Follows WCAG 2.1 AA and Microsoft Accessibility Checker standards
"""

import re
import sys
import argparse

import requests
from bs4 import BeautifulSoup
import pdfplumber
from PIL import Image
import pytesseract

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def mark_table_header_row(table):
    """Mark the first row of a table as a header row for accessibility"""
    tblPr = table._tblPr
    tblHeader = OxmlElement('w:tblHeader')
    table.rows[0]._element.get_or_add_trPr().append(tblHeader)


def create_menu_section(doc, category_heading, items):
    """
    Create an accessible menu section using headings instead of tables.
    items: list of dicts with 'name', 'description', 'price' keys

    Each item becomes a Heading 3 under the category (Heading 2) with a
    paragraph describing the item. The price is appended to the heading for
    improved clarity and ease of scanning.
    """
    # Add category heading
    doc.add_heading(category_heading, level=2)
    
    for item in items:
        # Combine name and price in heading
        name = item.get('name', '')
        price = item.get('price', '')
        heading_text = f"{name}"
        if price:
            heading_text += f" — {price}"
        item_heading = doc.add_heading(heading_text, level=3)
        
        # Paragraph for description only
        desc_text = item.get('description', '')
        if desc_text:
            para = doc.add_paragraph(desc_text)
        
        # small space after each item
        doc.add_paragraph()



# existing hardcoded data functions remain for reference but are no longer used by CLI

def parse_html_menu(url):
    """Fetch the URL and parse menu categories and items into a dict.

    Returns dict {category: [ {name,description,price}, ... ] }
    """
    resp = requests.get(url)
    resp.raise_for_status()
    soup = BeautifulSoup(resp.text, "html.parser")
    # try to extract page title
    h1 = soup.find('h1')
    page_title = h1.get_text(strip=True) if h1 else url

    menu = {}
    # categories appear as h3, items as h4 with price text nearby
    for cat in soup.find_all(['h2','h3']):
        cat_text = cat.get_text(strip=True)
        # skip non-menu headings
        if re.search(r'\$\d', cat_text):
            continue
        # collect following siblings until next same-level heading
        items = []
        for sib in cat.find_next_siblings():
            if sib.name in ['h2','h3']:
                break
            # look for item headings
            if sib.name in ['h4','h5']:
                name = sib.get_text(strip=True)
                # price might be in same element or next text
                price = ''
                # look for text nodes that look like price
                text = sib.get_text(separator=' ').strip()
                m = re.search(r'(\$?\d+\.?\d*)', text)
                if m:
                    price = m.group(1)
                # description may be in next paragraph
                desc = ''
                nxt = sib.find_next_sibling()
                if nxt and nxt.name == 'p':
                    desc = nxt.get_text(strip=True)
                items.append({'name':name,'description':desc,'price':price})
        if items:
            menu[cat_text] = items
    return page_title, menu


def parse_pdf_menu(path):
    """Basic PDF text extraction and simple menu parsing."""
    text = ''
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text += page.extract_text() + '\n'
    return _parse_text_to_menu(path, text)


def parse_image_menu(path):
    """Use OCR to extract text from an image and parse it into a menu."""
    img = Image.open(path)
    text = pytesseract.image_to_string(img)
    return _parse_text_to_menu(path, text)


def _parse_text_to_menu(source, text):
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    menu = {}
    current_cat = None
    for line in lines:
        if not re.search(r'\$?\d', line) and len(line) < 50:
            current_cat = line
            menu[current_cat] = []
            continue
        m = re.search(r'(.+?)\s+\$?(\d+\.?\d*)', line)
        if m and current_cat:
            name = m.group(1).strip()
            price = '$' + m.group(2)
            menu[current_cat].append({'name':name,'description':'','price':price})
            continue
        if current_cat and menu[current_cat]:
            if not re.search(r'\$?\d', line):
                menu[current_cat][-1]['description'] += ' ' + line
    return source, menu


def create_menu_section(doc, category_heading, items):
    """
    Create an accessible menu section using headings instead of tables.
            {"name": "#1 Breakfast", "description": "2 Eggs, Hashbrowns, Hotcakes or Toast & Coffee", "price": "$10.99"},
            {"name": "#2 The Blimp", "description": "2 Eggs, Hashbrowns, 2 Hotcakes or French Toast & Bacon, Sausage Links, or Ham", "price": "$14.79"},
            {"name": "#3 Avocado Toast with Feta & Tomato", "description": "with 2 Eggs (any style) and Fresh Fruit cup", "price": "$14.99"},
            {"name": "#4 Monte Cristo", "description": "Grilled Ham & Cheese between Golden French Toast with Hashbrowns", "price": "$10.99"},
            {"name": "#5 Akron Special", "description": "3 Eggs, Toast and Bacon, Sausage Links, or Ham", "price": "$9.79"},
            {"name": "#6 Healthy Favorite", "description": "Spinach Scrambler with Turkey links, Fresh Fruit & Toast", "price": "$13.29"},
            {"name": "#7 Eggs Benedict", "description": "Poached Eggs and Ham on English Muffin with Hollandaise Sauce and Hashbrowns. Sub salmon, crab or lobster +$3.00", "price": "$13.79"},
            {"name": "#8 Homemade Hash", "description": "with 2 Eggs and Toast", "price": "$11.79"},
        ],
        "AFR Favorites": [
            {"name": "Biscuits Spicy Sausage Gravy & Hashbrowns", "description": "2 Eggs, served with biscuits and spicy sausage gravy", "price": "$12.49"},
            {"name": "Hungarian Sausage & Two Eggs", "description": "Hashbrowns, Toast", "price": "$12.49"},
            {"name": "Turkey Sausage & Two Eggs", "description": "Hashbrowns, Toast", "price": "$12.49"},
            {"name": "Spicy Sausage Patty & Two Eggs", "description": "Hashbrowns, Toast", "price": "$12.49"},
            {"name": "Smoked Sausage & Two Eggs", "description": "Hashbrowns, Toast", "price": "$12.49"},
        ],
        "Omelets": [
            {"name": "Crab & Swiss", "description": "3 Egg Omelet with crab and Swiss cheese, served with hashbrowns and toast", "price": "$14.79"},
            {"name": "Spinach & Feta", "description": "3 Egg Omelet with spinach and feta cheese, served with hashbrowns and toast", "price": "$14.29"},
            {"name": "Gyro & Feta Omelet", "description": "with tomato & onion, served with hashbrowns and toast", "price": "$14.29"},
            {"name": "Western & Cheese", "description": "with ham, American cheese, green pepper, onion, served with hashbrowns and toast", "price": "$14.29"},
            {"name": "Mexican Omelet", "description": "with spicy sausage, cheddar, jalapeno, tomato, onion and side of salsa and sour cream", "price": "$16.99"},
            {"name": "Ham or Bacon or Sausage & Cheese", "description": "with American cheese, served with hashbrowns and toast", "price": "$14.29"},
            {"name": "Italian Sausage w/marinara & Mozzarella", "description": "served with hashbrowns and toast", "price": "$14.29"},
            {"name": "Chili Cheese Omelet", "description": "with American cheese and shredded cheddar on top, served with hashbrowns and toast", "price": "$14.29"},
            {"name": "Cheese Omelet", "description": "American, Swiss, or Cheddar, served with hashbrowns and toast", "price": "$12.29"},
            {"name": "Veggie Omelet", "description": "with Swiss Cheese, served with hashbrowns and toast", "price": "$13.79"},
            {"name": "Mushroom & Swiss Omelet", "description": "with Swiss Cheese, served with hashbrowns and toast", "price": "$14.29"},
        ],
        "Eggs": [
            {"name": "One egg, any style", "description": "Served with Hashbrowns and Toast", "price": "$6.89"},
            {"name": "One egg, any style w/ bacon, sausage links, or ham", "description": "Served with Hashbrowns and Toast", "price": "$8.39"},
            {"name": "Two eggs, any style", "description": "Served with Hashbrowns and Toast", "price": "$7.39"},
            {"name": "Two eggs, any style w/ bacon, sausage links, or ham", "description": "Served with Hashbrowns and Toast", "price": "$10.99"},
            {"name": "Three eggs, any style w/ bacon, sausage links, or ham", "description": "Served with Hashbrowns and Toast", "price": "$11.39"},
            {"name": "Two eggs, any style w/ Canadian Bacon", "description": "Served with Hashbrowns and Toast", "price": "$12.39"},
            {"name": "Gyro Breakfast", "description": "Two eggs any style, gyro meat, pita bread", "price": "$14.89"},
            {"name": "Chopped Steak and Eggs", "description": "Two eggs any style, 8oz chopped steak", "price": "$15.99"},
            {"name": "Sirloin Steak and Eggs", "description": "Two eggs any style, 6oz sirloin steak", "price": "$16.79"},
        ],
        "Breakfast Sandwiches": [
            {"name": "Egg and Cheese", "description": "2 eggs, American cheese on Kaiser bun with lettuce & tomato and chips", "price": "$8.59"},
            {"name": "Western Omelet Sandwich", "description": "2 eggs, Ham, Green Peppers, Onions, and American cheese on Kaiser bun", "price": "$9.49"},
            {"name": "Bacon, Ham, or Sausage Egg & Cheese Sandwich", "description": "2 eggs, American cheese on Kaiser bun with lettuce & tomato and chips", "price": "$9.49"},
        ],
        "Hot Cakes & French Toast": [
            {"name": "3 Blueberry or Chocolate Chip Hot Cakes", "description": "with bacon/sausage links/ham OR two eggs", "price": "$9.89 / $11.49"},
            {"name": "3 Buttermilk Hot Cakes OR 3 Texas French Toast", "description": "with bacon/sausage links/ham OR two eggs", "price": "$9.39 / $10.89"},
            {"name": "2 Blueberry or Chocolate Chip Hot Cakes", "description": "with bacon/sausage links/ham OR two eggs", "price": "$9.39 / $11.39"},
            {"name": "2 Buttermilk Hot Cakes OR 2 Texas French Toast", "description": "with bacon/sausage links/ham OR two eggs", "price": "$8.79 / $10.79"},
        ],
        "Waffles": [
            {"name": "Chicken N' Waffle", "description": "Three chicken tenders served on top of our Belgian Waffle", "price": "$14.59"},
            {"name": "Waffle Blimp", "description": "Two eggs, hashbrowns, waffle, w/ bacon, sausage links, or ham", "price": "$14.59"},
            {"name": "Blueberry Waffle", "description": "Blueberries, glazed with whipped cream", "price": "$9.99"},
            {"name": "Cinnamon Apple Waffle", "description": "Hot apples with whipped cream", "price": "$9.99"},
            {"name": "Peach Waffle", "description": "Peaches, glazed with whipped cream", "price": "$9.99"},
            {"name": "Strawberry Waffle", "description": "Strawberries with whipped cream", "price": "$9.99"},
            {"name": "Chocolate Chip OR Pecan Waffle", "description": "Chocolate chips or crushed pecans with whipped cream", "price": "$9.99"},
            {"name": "Belgian Waffle", "description": "Topped with powdered sugar", "price": "$7.99"},
        ],
        "Breads": [
            {"name": "Plain Bagel", "description": "with cream cheese", "price": "$5.39 / $5.79"},
            {"name": "English Muffin", "description": "", "price": "$4.49"},
            {"name": "Cinnamon Toast", "description": "Old fashioned buttered texas toast served with cinnamon sugar", "price": "$4.79"},
            {"name": "Muffin", "description": "Blueberry, Banana, or Cinnamon Crunch", "price": "$5.79"},
            {"name": "Toast", "description": "2 pieces of white, wheat, rye, or sourdough", "price": "$4.69"},
            {"name": "Pita Bread", "description": "", "price": "$3.79"},
        ],
        "Breakfast Sides": [
            {"name": "Egg", "description": "1 egg", "price": "$2.99"},
            {"name": "Oatmeal", "description": "served with milk and brown sugar", "price": "$6.69"},
            {"name": "Hashbrowns", "description": "", "price": "$5.29"},
            {"name": "Fresh Fruit", "description": "cup or bowl", "price": "$4.89 / $5.89"},
            {"name": "Spicy Sausage Patty", "description": "Local Canal Fulton Provisions", "price": "$4.99"},
            {"name": "Hungarian Sausage", "description": "Local Canal Fulton Provisions cut into 3 pieces", "price": "$4.99"},
            {"name": "Sausage Patties", "description": "2 patties", "price": "$4.89"},
            {"name": "Smoked Sausage", "description": "Cut into 3 pieces, local Canal Fulton Provisions Smoked Sausage", "price": "$4.89"},
            {"name": "Bacon or Sausage Links", "description": "3 pieces", "price": "$4.69"},
            {"name": "Turkey Sausage", "description": "3 pieces", "price": "$4.89"},
            {"name": "Ham", "description": "1 piece", "price": "$4.89"},
            {"name": "Canadian Bacon", "description": "3 pieces", "price": "$4.89"},
            {"name": "Gyro Meat", "description": "", "price": "$4.99"},
        ],
    }


def create_lunch_dinner_menu_data():
    """Return lunch and dinner menu data"""
    return {
        "Soups": [
            {"name": "Homemade Soup Special", "description": "ask your server for the daily specialty soup option", "price": "$3.79 cup / $4.69 bowl"},
            {"name": "Chicken Soup", "description": "", "price": "$3.79 cup / $4.69 bowl"},
            {"name": "Chili", "description": "supreme with onions, cheese, and sour cream", "price": "$4.99 cup / $6.49 bowl, $8.29 supreme"},
        ],
        "Akron's Famous Salads": [
            {"name": "Sirloin Steak Salad", "description": "tender 5oz steak, shredded cheddar & mozzarella cheese, green pepper, onion, egg, tomato, French fries and garlic bread", "price": "$15.29"},
            {"name": "Greek Salmon Salad", "description": "feta, black olives, egg, tomato, onion, pepperoncini, & pita", "price": "$14.99"},
            {"name": "Steak Salad", "description": "roast beef warmed, shredded cheddar & mozzarella cheese, green pepper, onion, egg, tomato, French fries & garlic bread", "price": "$12.99 small / $13.99 large"},
            {"name": "Julienne Salad", "description": "turkey, ham, roast beef, cheddar & mozzarella cheese, egg, tomato & garlic bread", "price": "$12.99 small / $13.99 large"},
            {"name": "Chef's Salad", "description": "turkey, ham, roast beef, cheddar & mozzarella cheese, egg, tomato & garlic bread", "price": "$12.99 small / $13.99 large"},
            {"name": "Greek Salad", "description": "feta cheese, onion, green pepper, black olives, tomato, egg, pepperoncini, house greek dressing & pita bread", "price": "$12.99 small / $13.99 large"},
            {"name": "Gyro Salad", "description": "gyro meat, onion, egg, tomato, French fries, & pita bread, mozzarella cheese", "price": "$12.99 small / $13.99 large"},
            {"name": "Grilled Chicken Breast Salad", "description": "grilled chicken breast, cheddar & mozzarella cheese, egg, tomato, & garlic bread", "price": "$12.99 small / $13.99 large"},
            {"name": "Cobb Salad", "description": "grilled chicken breast, bacon, mozzarella, cheddar & crumbled bleu cheese, egg, tomato & garlic bread", "price": "$12.99 small / $14.99 large"},
        ],
        "Akron's Famous Sandwiches": [
            {"name": "#1 Hot Corned Beef", "description": "on grilled rye with lettuce, tomato, Swiss cheese", "price": "$12.49"},
            {"name": "#2 Turkey Burger", "description": "on ciabatta roll with green pepper & pepper jack cheese", "price": "$12.49"},
            {"name": "#3 Beer Battered Cod Sandwich", "description": "with cheese & homemade tangy tarter sauce, lettuce & tomato", "price": "$12.99"},
            {"name": "#4 Roast Beef Melt", "description": "on sourdough with grilled onions, & cheddar cheese", "price": "$12.49"},
            {"name": "#5 Tuna or Chicken Salad", "description": "on toasted rye with lettuce, tomato, Swiss cheese", "price": "$12.49"},
            {"name": "#6 Turkey and Bacon Melt", "description": "on sourdough with cheddar cheese", "price": "$12.49"},
            {"name": "#7 Hamburger", "description": "with lettuce and tomato", "price": "$11.79"},
            {"name": "#8 Club Sandwich", "description": "on sourdough with turkey or ham and bacon, lettuce, tomato, cheese, & mayo", "price": "$12.49"},
            {"name": "#9 Lobster Swiss Bacon Melt", "description": "on sourdough with lettuce and tomato", "price": "$12.49"},
            {"name": "#10 Gyro Sandwich", "description": "gyro meat or grilled chicken with tomato, onion, & tzatziki sauce", "price": "$12.49"},
            {"name": "#11 Reuben Sandwich", "description": "BEST SELLER! on rye, grilled & piled high w/ corned beef or turkey along with sauerkraut and Swiss cheese", "price": "$13.29"},
            {"name": "#12 Sunrise Burger", "description": "with egg, cheese, bacon, lettuce & tomato", "price": "$12.99"},
            {"name": "#13 Chicken Parm Sandwich", "description": "with lettuce & tomato", "price": "$12.99"},
            {"name": "#14 Grilled Cheese Sandwich", "description": "on texas toast", "price": "$8.79"},
            {"name": "#15 Grilled Ham or Bacon & Cheese", "description": "on texas toast", "price": "$11.29"},
            {"name": "#16 BLT Sandwich", "description": "on grilled sourdough with mayo", "price": "$11.29"},
            {"name": "#17 Mushroom Burger", "description": "with lettuce, tomato, Swiss cheese", "price": "$12.79"},
            {"name": "#18 Bacon Burger", "description": "with lettuce, tomato, American cheese", "price": "$12.79"},
            {"name": "#19 Patty Melt", "description": "on rye with Swiss cheese & grilled onions", "price": "$12.79"},
            {"name": "#20 Grilled Tuna Melt", "description": "on texas toast with lettuce, tomato, & cheese", "price": "$12.79"},
            {"name": "#21 Meatball Sub", "description": "with cheese", "price": "$12.79"},
            {"name": "#22 Baja Burger", "description": "Spicy - with pepper jack cheese, jalapeños & green peppers", "price": "$12.79"},
            {"name": "#23 Black & Bleu Burger", "description": "with melted blue cheese crumbles, lettuce & tomatoes", "price": "$12.79"},
        ],
        "Sub Sandwiches": [
            {"name": "Strip Steak Sub", "description": "with grilled onions/mushrooms, lettuce, tomato & Swiss cheese, served with French fries", "price": "$16.99"},
            {"name": "Philly Steak Sub", "description": "with grilled onions, mushrooms & Swiss, served with French fries", "price": "$15.99"},
            {"name": "Italian Sausage Sub", "description": "with grilled green peppers & onions, served with French fries", "price": "$15.99"},
        ],
        "Special Sandwiches": [
            {"name": "Hot Roast Beef", "description": "open face with mashed potatoes & gravy", "price": "$14.99"},
            {"name": "Chicken Quesadilla", "description": "with chicken, cheese, green peppers, onions served with French fries", "price": "$15.99"},
            {"name": "Salmon BLT", "description": "with mayo served with sweet potato fries", "price": "$16.99"},
        ],
        "Chicken Sandwiches": [
            {"name": "Grilled Chicken Breast", "description": "loaded with Bacon/Cheddar, served with French fries", "price": "$14.99"},
            {"name": "Grilled Chicken Breast", "description": "loaded with Mushroom/Swiss, served with French fries", "price": "$14.99"},
            {"name": "Chicken Avocado BLT", "description": "served on ciabatta bread with French fries", "price": "$15.79"},
        ],
        "Healthy Burgers, Wraps and Pitas": [
            {"name": "Black Bean Burger", "description": "lettuce, tomato served with pepper jack cheese, grilled peppers & sweet potato fries", "price": "$14.99"},
            {"name": "Veggie Burger", "description": "lettuce, tomato served with pepper jack cheese, grilled peppers & sweet potato fries", "price": "$14.99"},
            {"name": "BLT Wrap/Pita", "description": "bacon, lettuce, tomato served with French fries", "price": "$14.99"},
            {"name": "Chicken Tender Wrap/Pita", "description": "with cheese, lettuce, tomato served with French fries", "price": "$14.99"},
            {"name": "Tuna Salad Wrap/Pita", "description": "with lettuce, tomato served with French fries", "price": "$14.99"},
            {"name": "Chicken Salad Wrap/Pita", "description": "with lettuce, tomato served with French fries", "price": "$14.99"},
            {"name": "Philly Steak Wrap/Pita", "description": "grilled onions/mushrooms/swiss served with French fries", "price": "$14.99"},
        ],
        "Akron's Specials": [
            {"name": "Weight Watchers Special", "description": "8oz chopped steak or grilled chicken breast and melted Swiss cheese on bed of crisp green lettuce, cottage cheese, hard-boiled egg, sliced tomatoes, and pita bread", "price": "$14.79"},
            {"name": "French Dip Platter", "description": "warmed roast beef with au jus, piled high on grilled French bread served with sweet potato fries and coleslaw", "price": "$14.79"},
            {"name": "Chicken or Tuna Salad Plate", "description": "large serving garnished with fruit and cottage cheese served with tomato slices on bed of crisp green lettuce and pita bread", "price": "$14.79"},
            {"name": "Gyro Platter", "description": "gyro meat, sliced tomato, onion, lettuce & pita bread served with French fries", "price": "$14.79"},
        ],
        "Entrees": [
            {"name": "Strip Steak", "description": "with onion & mushrooms, served with vegetables, potato, and choice of soup or tossed salad", "price": "$17.99"},
            {"name": "Country Fried Steak", "description": "with spicy sausage gravy, served with vegetables, potato, and choice of soup or tossed salad", "price": "$16.99"},
            {"name": "Sirloin Steak Dinner", "description": "with onions and mushrooms, served with vegetables, potato, and choice of soup or tossed salad", "price": "$16.99"},
            {"name": "Grilled Chicken Breast", "description": "with gravy, served with vegetables, potato, and choice of soup or tossed salad", "price": "$16.99"},
            {"name": "Grilled Ham Steak & Pineapple", "description": "with gravy, served with vegetables, potato, and choice of soup or tossed salad", "price": "$16.99"},
            {"name": "Smothered Chicken", "description": "with mushrooms, onion, & Swiss cheese, served with vegetables, potato, and choice of soup or tossed salad", "price": "$16.99"},
            {"name": "Liver & Onions", "description": "beef liver and sautéed onions, served with vegetables, potato, and choice of soup or tossed salad", "price": "$16.99"},
            {"name": "Chopped Steak", "description": "with mushrooms and onions, served with vegetables, potato, and choice of soup or tossed salad", "price": "$16.99"},
            {"name": "Chicken Strip Dinner", "description": "4 tender chicken strips, served with vegetables, potato, and choice of soup or tossed salad", "price": "$16.99"},
        ],
        "Seafood Specials": [
            {"name": "Fried Clam Strips Dinner", "description": "crispy golden clam strips fried to perfection, served with cocktail sauce, vegetables, potato, and choice of soup or tossed salad", "price": "$15.29"},
            {"name": "Beer Battered Cod Dinner", "description": "Atlantic cod dipped in crisp beer batter, served golden brown with homemade tarter sauce, vegetables, potato, and choice of soup or tossed salad", "price": "$15.99"},
            {"name": "Fantail Shrimp Dinner", "description": "lightly breaded fantail shrimp fried to golden crisp served with cocktail sauce, vegetables, potato, and choice of soup or tossed salad", "price": "$15.99"},
        ],
        "Pasta": [
            {"name": "Spaghetti OR Penne with Marinara", "description": "Served with garlic bread and your choice of soup or tossed salad", "price": "$12.89"},
            {"name": "Spaghetti OR Penne with Marinara & Meatballs", "description": "Served with garlic bread and your choice of soup or tossed salad", "price": "$13.99"},
            {"name": "Italian Sausage Pasta Bowl", "description": "Served with garlic bread and your choice of soup or tossed salad", "price": "$13.99"},
            {"name": "Chicken Parmesan", "description": "served with spaghetti or penne and marinara, garlic bread and your choice of soup or tossed salad", "price": "$14.99"},
        ],
        "Sides": [
            {"name": "French Fries", "description": "", "price": "$5.29"},
            {"name": "Onion Rings", "description": "", "price": "$5.89"},
            {"name": "Sweet Potato Fries", "description": "", "price": "$5.39"},
            {"name": "Applesauce", "description": "", "price": "$3.29"},
            {"name": "Cottage Cheese", "description": "", "price": "$3.29"},
            {"name": "Coleslaw", "description": "", "price": "$3.29"},
            {"name": "Vegetables", "description": "", "price": "$3.29"},
            {"name": "Tossed Salad", "description": "", "price": "$5.29"},
            {"name": "Garlic Bread", "description": "", "price": "$3.59"},
            {"name": "Pita Bread", "description": "", "price": "$3.59"},
        ],
        "Lunch & Dinner Beverages": [
            {"name": "Coffee or Tea", "description": "", "price": "$3.39"},
            {"name": "Juice", "description": "orange, apple, cranberry, grapefruit, tomato", "price": "$3.59"},
            {"name": "Iced Tea", "description": "", "price": "$3.39"},
            {"name": "Pink Lemonade", "description": "", "price": "$3.39"},
            {"name": "Pepsi Products", "description": "", "price": "$3.39"},
            {"name": "Milk", "description": "white or chocolate", "price": "$3.59"},
            {"name": "Hot Chocolate", "description": "with whipped cream", "price": "$4.29"},
            {"name": "Cappuccino", "description": "with whipped cream", "price": "$4.79"},
        ],
        "Desserts": [
            {"name": "Cobbler", "description": "served with ice cream", "price": "$4.99"},
            {"name": "Boston Cream Pie", "description": "", "price": "$4.99"},
        ],
    }




def build_doc_from_menu(menu_dict, title, output_path):
    """Generate a Word document from a menu dictionary.

    menu_dict: {category: [ {name,description,price}, ... ] }
    title: string for document title
    output_path: path to save .docx
    """
    doc = Document()
    # properties
    doc.core_properties.title = title
    doc.core_properties.author = "Menu Generator"
    doc.core_properties.language_id = 1033
    
    # main heading
    doc.add_paragraph(title, style='Title')
    doc.add_paragraph()  # space
    
    for category, items in menu_dict.items():
        doc.add_heading(category, level=1)
        for item in items:
            # combine name and price
            name = item.get('name','')
            price = item.get('price','')
            heading_text = name
            if price:
                heading_text += f" — {price}"
            doc.add_heading(heading_text, level=2)
            desc = item.get('description','')
            if desc:
                doc.add_paragraph(desc)
            doc.add_paragraph()
    
    # footer info
    doc.add_page_break()
    doc.add_heading("Notes", level=1)
    doc.add_paragraph("Consuming raw or undercooked meats, poultry, seafood or eggs may increase your risk of food borne illness.")
    
    doc.save(output_path)
    print(f"Saved to {output_path}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Accessible menu to Word converter")
    parser.add_argument('--url', help='URL of the menu page')
    parser.add_argument('--pdf', help='Path to PDF menu file')
    parser.add_argument('--image', help='Path to image file containing menu (JPG/PNG)')
    parser.add_argument('--output','-o', help='Output docx file', default='menu.docx')
    args = parser.parse_args()
    if args.url:
        title, menu = parse_html_menu(args.url)
    elif args.pdf:
        title, menu = parse_pdf_menu(args.pdf)
    elif args.image:
        title, menu = parse_image_menu(args.image)
    else:
        parser.error("one of --url, --pdf or --image required")
    build_doc_from_menu(menu, title, args.output)
